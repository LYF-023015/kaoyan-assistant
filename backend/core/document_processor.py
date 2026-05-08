"""
文档处理器：统一加载、分块、元数据注入
"""
import os
import re
from pathlib import Path
from typing import List
from backend.core.pdf_parser import PDFParser, Document


class DocumentProcessor:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 150):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.pdf_parser = PDFParser()
        self.supported_exts = {".txt", ".md", ".pdf", ".docx"}

    def process_directory(self, directory: str, subject: str = "") -> List[Document]:
        """
        处理目录下所有支持的文件
        """
        path = Path(directory)
        all_docs = []
        for file_path in path.rglob("*"):
            if file_path.suffix.lower() in self.supported_exts:
                docs = self.process_file(str(file_path), subject)
                all_docs.extend(docs)
        return all_docs

    def process_file(self, file_path: str, subject: str = "") -> List[Document]:
        """
        处理单个文件
        """
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self.pdf_parser.parse(file_path, subject)
        elif ext == ".txt":
            return self._process_text(file_path, subject)
        elif ext == ".md":
            return self._process_markdown(file_path, subject)
        elif ext == ".docx":
            return self._process_docx(file_path, subject)
        return []

    def _process_text(self, file_path: str, subject: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        chunks = self._recursive_split(text)
        return [
            Document(
                content=chunk,
                metadata={
                    "source": Path(file_path).name,
                    "subject": subject,
                    "file_type": "txt",
                    "chunk_index": i,
                }
            )
            for i, chunk in enumerate(chunks)
        ]

    def _process_markdown(self, file_path: str, subject: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        # 简单去除部分 Markdown 语法
        text = re.sub(r'#{1,6}\s*', '', text)
        text = re.sub(r'\*\*?(.*?)\*\*?', r'\1', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        chunks = self._recursive_split(text)
        return [
            Document(
                content=chunk,
                metadata={
                    "source": Path(file_path).name,
                    "subject": subject,
                    "file_type": "md",
                    "chunk_index": i,
                }
            )
            for i, chunk in enumerate(chunks)
        ]

    def _process_docx(self, file_path: str, subject: str) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return []

        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        chunks = self._recursive_split(full_text)
        return [
            Document(
                content=chunk,
                metadata={
                    "source": Path(file_path).name,
                    "subject": subject,
                    "file_type": "docx",
                    "chunk_index": i,
                }
            )
            for i, chunk in enumerate(chunks)
        ]

    def _recursive_split(self, text: str) -> List[str]:
        """
        递归字符文本分割，对中文更友好
        """
        separators = ["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]
        chunks = []
        self._split_text_recursive(text, separators, 0, chunks)
        return chunks

    def _split_text_recursive(self, text: str, separators: List[str], sep_idx: int, chunks: List[str]):
        if not text.strip():
            return
        if len(text) <= self.chunk_size:
            chunks.append(text.strip())
            return

        if sep_idx >= len(separators):
            # 强制切分
            for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
                chunk = text[i:i + self.chunk_size].strip()
                if chunk:
                    chunks.append(chunk)
            return

        sep = separators[sep_idx]
        if sep:
            parts = text.split(sep)
        else:
            # 空分隔符，按字符切分
            parts = list(text)

        current = ""
        for part in parts:
            candidate = current + sep + part if current and sep else part
            if len(candidate) <= self.chunk_size:
                current = candidate
            else:
                if current:
                    self._split_text_recursive(current, separators, sep_idx + 1, chunks)
                current = part
        if current:
            self._split_text_recursive(current, separators, sep_idx + 1, chunks)
